from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_UNDERLINE
import difflib


def highlight_changes(original_doc_path, changed_doc_path, output_doc_path):

    original_doc = Document(original_doc_path)
    changed_doc = Document(changed_doc_path)

    output_doc = Document()

    for original_para, changed_para in zip(original_doc.paragraphs, changed_doc.paragraphs):
        output_para = output_doc.add_paragraph()

        d = difflib.Differ()
        diff = list(d.compare(original_para.text.split(), changed_para.text.split()))

        # Add the changes to the output document in red
        for word in diff:
            if word.startswith("+ "):  # This word is added in the changed document
                run = output_para.add_run(f" {word[2:]}")
                run.font.color.rgb = RGBColor(0, 255, 0)  # Red color
            elif word.startswith("- "):
                run = output_para.add_run(f" {word[2:]}")
                run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                run = output_para.add_run(f" {word[2:]}")
        # Add a new line after each paragraph comparison
        output_doc.add_paragraph()

    # Save the output document
    output_doc.save(output_doc_path)

if __name__ == "__main__":
    # Paths to the original, changed, and output documents
    original_doc_path = "doc_original.docx"
    changed_doc_path = "doc_changed.docx"
    output_doc_path = "highlighted_changes.docx"

    # Run the function to highlight changes
    highlight_changes(original_doc_path, changed_doc_path, output_doc_path)
